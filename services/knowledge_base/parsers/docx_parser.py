"""
DOCX parser — extracts text with heading structure.

Uses python-docx to extract paragraphs grouped by headings.
Returns list of ParsedSection with text and metadata.
"""

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    text: str
    heading: str = ""


def parse_docx(file_path: Path) -> list[ParsedSection]:
    """
    Parse a DOCX file into sections grouped by headings.

    Each heading starts a new section. Text without a preceding heading
    goes into a section with empty heading.
    """
    try:
        doc = Document(str(file_path))
    except Exception as e:
        logger.error("Failed to parse DOCX %s: %s", file_path.name, e)
        return []

    sections: list[ParsedSection] = []
    current_heading = ""
    current_paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a heading
        style_name = (para.style.name or "").lower()
        if "heading" in style_name or "заголовок" in style_name:
            # Save previous section
            if current_paragraphs:
                full_text = "\n".join(current_paragraphs)
                if full_text.strip():
                    sections.append(ParsedSection(
                        text=full_text,
                        heading=current_heading,
                    ))
                current_paragraphs = []
            current_heading = text
        else:
            current_paragraphs.append(text)

    # Save last section
    if current_paragraphs:
        full_text = "\n".join(current_paragraphs)
        if full_text.strip():
            sections.append(ParsedSection(
                text=full_text,
                heading=current_heading,
            ))

    # If no sections found (no headings in doc), treat whole doc as one section
    if not sections:
        all_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        if all_text:
            sections.append(ParsedSection(text=all_text))

    # Also extract text from tables
    table_texts = []
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            table_texts.append("\n".join(rows_text))

    if table_texts:
        sections.append(ParsedSection(
            text="\n\n".join(table_texts),
            heading="Таблицы",
        ))

    total_chars = sum(len(s.text) for s in sections)
    logger.info(
        "Parsed DOCX %s: %d sections, %d chars",
        file_path.name, len(sections), total_chars,
    )
    return sections
